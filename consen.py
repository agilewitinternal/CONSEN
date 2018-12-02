import pandas as pd
import tkinter as tk
import xlrd
import docx
from tkinter import filedialog
from tkinter import messagebox


def load_excel():
    messagebox.showinfo("Excel File", "Select EXCEL File to Load")
    fileName = filedialog.askopenfilename()
    wb = xlrd.open_workbook(fileName)
    return wb

def File_Load():
    messagebox.showinfo("Doc File", "Select Docx File to Modify")
    fileName = filedialog.askopenfilename()
    return fileName

def PathFile():
    master = Tk()
    Label(master, text="Input: ").grid(row=0, sticky=W)
    PathFile = entry.get()
    entry = Entry(master)
    entry.grid(row=0, column=1)

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

workbook = load_excel()
Doc_Path = File_Load()
sheet = workbook.sheet_by_index(0)
sheet.cell_value(0, 0)

for i in range(sheet.nrows):
    if i == 0:
        continue
    else:
        sheet_data = sheet.row_values(i)
        wordfile = getText(Doc_Path).format(sheet_data[0],sheet_data[1])
        wr1 = docx.Document(r'E:\TUTORIALS_AND_CODINGS\Python_Coding_practice\CONSEN\Wor.docx')
        print(wordfile)
        wr1.add_paragraph(wordfile)
        wr1.save(f'E:\TUTORIALS_AND_CODINGS\Python_Coding_practice\CONSEN\Created_Words\{sheet_data[0]}.docx')
        print('Created')